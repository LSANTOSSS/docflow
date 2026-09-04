from pathlib import Path
from typing import Any
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docflow.parser import Block


def _heading_size(base_size: float, level: int) -> float:
    return max(base_size - (level - 1) * 1.3, 12.0)


def _apply_document_config(document: Document, config: dict[str, Any]) -> None:
    metadata = config["document"]
    properties = document.core_properties
    for field in ("title", "author", "subject", "keywords"):
        value = metadata.get(field)
        if value:
            setattr(properties, field, str(value))

    styles = config["styles"]
    body = styles["body"]
    heading = styles["heading"]
    normal = document.styles["Normal"]
    normal.font.name = body["font"]
    normal.font.size = Pt(float(body["size"]))
    heading_base_size = float(heading.get("size") or 18)
    for level in range(1, 7):
        heading_style = document.styles[f"Heading {level}"]
        heading_style.font.name = heading["font"]
        heading_style.font.size = Pt(_heading_size(heading_base_size, level))

    for section in document.sections:
        if metadata.get("header"):
            section.header.paragraphs[0].text = str(metadata["header"])
        if metadata.get("footer"):
            section.footer.paragraphs[0].text = str(metadata["footer"])


def _add_code_block(document, block, code_style):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    if block.language:
        label = paragraph.add_run(f"{block.language}\n")
        label.bold = True
        label.font.name = code_style["font"]
        label.font.size = Pt(float(code_style["size"]))
    run = paragraph.add_run(block.text)
    run.font.name = code_style["font"]
    run.font.size = Pt(float(code_style["size"]))


def _add_table(document, block):
    if not block.rows:
        return
    table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(block.rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _create_document(config):
    template_path = config.get("template", {}).get("path")
    return Document(str(template_path)) if template_path else Document()


def _add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Atualize o campo para gerar o sumário."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def export_docx(blocks: list[Block], output: Path, config: dict[str, Any] | None = None) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    if config is None:
        from docflow.config import DEFAULT_CONFIG
        config = DEFAULT_CONFIG

    document = _create_document(config)
    _apply_document_config(document, config)

    if config.get("document", {}).get("toc"):
        _add_toc(document)

    code_style = config["styles"]["code"]
    for block in blocks:
        if block.kind == "heading":
            document.add_heading(block.text, level=block.level or 1)
        elif block.kind == "unordered_list":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "ordered_list":
            document.add_paragraph(block.text, style="List Number")
        elif block.kind == "code":
            _add_code_block(document, block, code_style)
        elif block.kind == "table":
            _add_table(document, block)
        else:
            document.add_paragraph(block.text)

    document.save(output)
    return output
