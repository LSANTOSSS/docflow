from pathlib import Path
import pytest
from docx import Document
from docflow.cli import run_export
from docflow.config import load_config
from docflow.parser import parse_markdown
from docflow.validation import validate_structure

def test_preset_applies_and_can_be_overridden(tmp_path: Path):
    cfg=tmp_path/"docflow.yaml"
    cfg.write_text("template:\n  preset: report\nstyles:\n  body:\n    size: 12\n",encoding="utf-8")
    config=load_config(cfg)
    assert config["styles"]["body"]["font"]=="Calibri"
    assert config["styles"]["body"]["size"]==12
    assert config["structure"]["required_headings"]==["Resumo"]

def test_unknown_preset_is_rejected(tmp_path: Path):
    cfg=tmp_path/"docflow.yaml"; cfg.write_text("template:\n  preset: nope\n",encoding="utf-8")
    with pytest.raises(ValueError,match="Preset desconhecido"): load_config(cfg)

def test_structure_validation():
    config={"structure":{"required_headings":["Resumo","Conclusão"]}}
    with pytest.raises(ValueError,match="Resumo, Conclusão"):
        validate_structure(parse_markdown("# Introdução\n"),config)
    validate_structure(parse_markdown("# Resumo\n\n# Conclusão\n"),config)

def test_export_uses_reference_template(tmp_path: Path):
    template=tmp_path/"reference.docx"; base=Document(); base.add_paragraph("Base pública"); base.save(template)
    source=tmp_path/"sample.md"; source.write_text("# Documento\n\nConteúdo.\n",encoding="utf-8")
    cfg=tmp_path/"docflow.yaml"; cfg.write_text("template:\n  path: reference.docx\n",encoding="utf-8")
    output=tmp_path/"out.docx"; run_export(source,output,cfg)
    doc=Document(output); text="\n".join(p.text for p in doc.paragraphs)
    assert "Base pública" in text and "Documento" in text and "Conteúdo." in text

def test_export_without_config_remains_compatible(tmp_path: Path):
    source=tmp_path/"sample.md"; source.write_text("# Documento\n\nTexto.\n",encoding="utf-8")
    output=tmp_path/"out.docx"; run_export(source,output)
    assert output.exists() and output.stat().st_size>0
