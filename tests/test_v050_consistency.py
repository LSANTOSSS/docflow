from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from docflow.cli import run_export
from docflow.config import load_config
from docflow.exporters.pdf import _styles


class TextCollector(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data.strip())


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            if item.text:
                parts.append(item.text)
        elif isinstance(item, Table):
            for row in item.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(parts)


def _html_text(path: Path) -> str:
    parser = TextCollector()
    parser.feed(path.read_text(encoding="utf-8"))
    return "\n".join(parser.parts)


def _pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def test_common_structure_is_preserved_across_formats(tmp_path: Path):
    source = tmp_path / "consistency.md"
    source.write_text(
        """# Consistency Document

Common paragraph.

- Bullet item

1. First ordered item
2. Second ordered item

| Field | Value |
| --- | --- |
| status | consistent |

```python
print("docflow")
```
""",
        encoding="utf-8",
    )

    outputs = {
        "docx": tmp_path / "consistency.docx",
        "html": tmp_path / "consistency.html",
        "pdf": tmp_path / "consistency.pdf",
    }
    for output in outputs.values():
        run_export(source, output)

    extracted = {
        "docx": _docx_text(outputs["docx"]),
        "html": _html_text(outputs["html"]),
        "pdf": _pdf_text(outputs["pdf"]),
    }
    expected = [
        "Consistency Document",
        "Common paragraph.",
        "Bullet item",
        "First ordered item",
        "Second ordered item",
        "Field",
        "Value",
        "status",
        "consistent",
        'print("docflow")',
    ]

    for format_name, text in extracted.items():
        for token in expected:
            assert token in text, f"{token!r} ausente em {format_name}"
        positions = [text.index(token) for token in expected]
        assert positions == sorted(positions), f"ordem estrutural divergente em {format_name}"

    assert "python" in extracted["docx"]
    assert "python" in extracted["pdf"]
    assert 'class="language-python"' in outputs["html"].read_text(encoding="utf-8")


def test_ordered_lists_keep_sequence_in_html_and_pdf(tmp_path: Path):
    source = tmp_path / "ordered.md"
    source.write_text(
        """# Ordered

1. First item
2. Second item
3. Third item
""",
        encoding="utf-8",
    )
    html_output = tmp_path / "ordered.html"
    pdf_output = tmp_path / "ordered.pdf"

    run_export(source, html_output)
    run_export(source, pdf_output)

    html = html_output.read_text(encoding="utf-8")
    pdf = _pdf_text(pdf_output)

    assert "<ol><li>First item</li><li>Second item</li><li>Third item</li></ol>" in html
    assert "1. First item" in pdf
    assert "2. Second item" in pdf
    assert "3. Third item" in pdf


def test_docx_uses_numbered_list_style_for_ordered_items(tmp_path: Path):
    source = tmp_path / "ordered.md"
    output = tmp_path / "ordered.docx"
    source.write_text("1. First item\n2. Second item\n", encoding="utf-8")

    run_export(source, output)

    document = Document(output)
    ordered = [paragraph for paragraph in document.paragraphs if paragraph.text]
    assert [paragraph.text for paragraph in ordered] == ["First item", "Second item"]
    assert all(paragraph.style.name == "List Number" for paragraph in ordered)


def test_typography_configuration_is_shared_across_exporters(tmp_path: Path):
    source = tmp_path / "typography.md"
    config_path = tmp_path / "docflow.yaml"
    source.write_text("# Title\n\n## Section\n\nText.\n\n```python\nprint('x')\n```\n", encoding="utf-8")
    config_path.write_text(
        "styles:\n"
        "  body:\n    font: Times New Roman\n    size: 12\n"
        "  heading:\n    font: Times New Roman\n    size: 20\n"
        "  code:\n    font: Courier New\n    size: 10\n",
        encoding="utf-8",
    )

    docx_output = tmp_path / "typography.docx"
    html_output = tmp_path / "typography.html"
    pdf_output = tmp_path / "typography.pdf"
    for output in (docx_output, html_output, pdf_output):
        run_export(source, output, config_path)

    document = Document(docx_output)
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert document.styles["Normal"].font.size.pt == 12
    assert document.styles["Heading 1"].font.name == "Times New Roman"
    assert document.styles["Heading 1"].font.size.pt == 20
    assert round(document.styles["Heading 2"].font.size.pt, 1) == 18.7

    html = html_output.read_text(encoding="utf-8")
    assert "font-family: Times New Roman, sans-serif; font-size: 12pt" in html
    assert "h1 { font-family: Times New Roman, sans-serif; font-size: 20pt" in html
    assert "h2 { font-family: Times New Roman, sans-serif; font-size: 18.7pt" in html
    assert "font-family: Courier New, monospace; font-size: 10pt" in html

    config = load_config(config_path)
    body, headings, code, _ = _styles(config)
    assert body.fontName == "Times-Roman"
    assert body.fontSize == 12
    assert headings[1].fontName == "Times-Roman"
    assert headings[1].fontSize == 20
    assert round(headings[2].fontSize, 1) == 18.7
    assert code.fontName == "Courier"
    assert code.fontSize == 10
