from pathlib import Path

from docx import Document

from docflow.cli import run_export


def test_export_creates_docx(tmp_path: Path):
    source = tmp_path / "sample.md"
    output = tmp_path / "sample.docx"
    source.write_text("# Documento\n\nConteúdo de teste.\n", encoding="utf-8")

    generated = run_export(source, output)

    assert generated == output
    assert output.exists()
    assert output.stat().st_size > 0


def test_export_supports_table_code_header_and_footer(tmp_path: Path):
    source = tmp_path / "sample.md"
    output = tmp_path / "sample.docx"
    config = tmp_path / "docflow.yaml"
    source.write_text(
        """# Documento

| Campo | Valor |
| --- | --- |
| versão | 0.2.0 |

```python
print("docflow")
```
""",
        encoding="utf-8",
    )
    config.write_text(
        "document:\n  header: DocFlow Portfolio\n  footer: v0.2.0\n",
        encoding="utf-8",
    )

    run_export(source, output, config)

    document = Document(output)
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "0.2.0"
    assert document.sections[0].header.paragraphs[0].text == "DocFlow Portfolio"
    assert document.sections[0].footer.paragraphs[0].text == "v0.2.0"
    assert "python" in "\n".join(p.text for p in document.paragraphs)
    assert 'print("docflow")' in "\n".join(p.text for p in document.paragraphs)
